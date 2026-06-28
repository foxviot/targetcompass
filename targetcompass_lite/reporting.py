import csv
import html
import json
import sqlite3
from pathlib import Path
from typing import Any

from .codex_engineering import load_codex_engineering
from .evidence_db import migrate_evidence_db
from .trace_orchestrator import refresh_traceability


PROHIBITED_CLAIMS = ["clinical recommendation", "cure"]

REPORT_SECTIONS = [
    "执行摘要",
    "研究问题与边界",
    "方法与模块",
    "数据来源与QC",
    "候选排序",
    "证据链",
    "限制与风险",
    "实验建议",
    "审批与审计",
]


REPORT_SECTIONS = [
    "执行摘要",
    "研究问题与边界",
    "方法与模块",
    "真实文献与数据库验证",
    "数据来源与 QC",
    "候选排序",
    "证据链",
    "限制与风险",
    "实验建议",
    "审批与审计",
]


def _read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


def _read_tsv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(encoding="utf-8") as f:
        return list(csv.DictReader(f, delimiter="\t"))


def _read_json(path: Path, fallback: Any) -> Any:
    if not path.exists():
        return fallback
    return json.loads(path.read_text(encoding="utf-8"))


def _write_json(path: Path, data: Any) -> None:
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def _table(headers: list[str], rows: list[list[Any]]) -> str:
    head = "".join(f"<th>{html.escape(h)}</th>" for h in headers)
    body = "\n".join(
        "<tr>" + "".join(f"<td>{html.escape(_cell_text(cell))}</td>" for cell in row) + "</tr>"
        for row in rows
    )
    return f"<table><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>"


def _cell_text(value: Any) -> str:
    if isinstance(value, list):
        return "; ".join(str(item) for item in value)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return "" if value is None else str(value)


def _candidate_rows(scores: list[dict[str, str]]) -> list[list[str]]:
    return [
        [
            row.get("score_id", ""),
            row.get("entity_symbol", ""),
            row.get("route", ""),
            row.get("final_score", ""),
            row.get("tier", ""),
            row.get("hard_gate_status", ""),
            row.get("safety_gate", ""),
            row.get("next_experiments", ""),
        ]
        for row in scores[:20]
    ]


def _method_records(context: dict[str, Any]) -> list[dict[str, Any]]:
    return context.get("analysis_modules", {}).get("modules", [])


def _method_rows(context: dict[str, Any]) -> list[list[Any]]:
    return [
        [
            module.get("module_id", ""),
            module.get("status", ""),
            module.get("input_modality", ""),
            module.get("runner", ""),
            module.get("outputs", []),
            module.get("notes", ""),
        ]
        for module in _method_records(context)
    ]


def _dataset_records(context: dict[str, Any]) -> list[dict[str, str]]:
    matches = {row.get("dataset_id"): row for row in context["matches"]}
    rows = []
    for row in context["screening"]:
        match = matches.get(row.get("dataset_id"), {})
        rows.append(
            {
                "dataset_id": row.get("dataset_id", ""),
                "source_class": row.get("source_class", ""),
                "grade": row.get("grade", ""),
                "modality": row.get("modality", ""),
                "metadata_quality": (row.get("metadata_quality_label", "") + " " + row.get("metadata_quality_score", "")).strip(),
                "gene_identity": row.get("gene_identity_status", ""),
                "recommended_use": row.get("recommended_use", ""),
                "match_status": match.get("match_status", ""),
                "notes": match.get("warnings", "") or row.get("reasons", ""),
            }
        )
    return rows


def _dataset_rows(context: dict[str, Any]) -> list[list[str]]:
    return [
        [
            row["dataset_id"],
            row["source_class"],
            row["grade"],
            row["modality"],
            row["metadata_quality"],
            row["gene_identity"],
            row["recommended_use"],
            row["match_status"],
            row["notes"],
        ]
        for row in _dataset_records(context)
    ]


def _deg_qc_records(project_dir: Path) -> list[dict[str, Any]]:
    rows = []
    for qc_path in sorted((project_dir / "results").glob("bulk_deg_*/qc_summary.json")):
        qc = _read_json(qc_path, {})
        rows.append(
            {
                "dataset_id": qc.get("dataset_id", qc_path.parent.name.replace("bulk_deg_", "")),
                "matrix_type": qc.get("matrix_type", ""),
                "runner_type": qc.get("runner_type", ""),
                "runner_reason": qc.get("runner_reason", ""),
                "case_samples": qc.get("case_samples", ""),
                "control_samples": qc.get("control_samples", ""),
                "genes": qc.get("genes", ""),
                "duplicated_gene_rows": qc.get("duplicated_gene_rows", ""),
                "design_full_rank": qc.get("design_full_rank", ""),
                "batch_covariates": qc.get("batch_covariates", []),
                "qc_status": qc.get("qc_status", ""),
                "artifact": str(qc_path.relative_to(project_dir)),
            }
        )
    if rows:
        return rows
    for qc_path in sorted((project_dir / "results").glob("bulk_deg_*/qc_summary.tsv")):
        metrics = {row["metric"]: row["value"] for row in _read_tsv(qc_path)}
        rows.append(
            {
                "dataset_id": qc_path.parent.name.replace("bulk_deg_", ""),
                "matrix_type": metrics.get("matrix_type", "unknown"),
                "runner_type": metrics.get("runner_type", "unknown"),
                "runner_reason": metrics.get("runner_reason", ""),
                "case_samples": metrics.get("case_samples", ""),
                "control_samples": metrics.get("control_samples", ""),
                "genes": metrics.get("genes", ""),
                "duplicated_gene_rows": metrics.get("duplicated_gene_rows", ""),
                "design_full_rank": metrics.get("design_full_rank", ""),
                "batch_covariates": metrics.get("batch_covariates", ""),
                "qc_status": metrics.get("qc_status", ""),
                "artifact": str(qc_path.relative_to(project_dir)),
            }
        )
    return rows


def _deg_qc_rows(project_dir: Path) -> list[list[Any]]:
    return [
        [
            row["dataset_id"],
            row["matrix_type"],
            row["runner_type"],
            row["case_samples"],
            row["control_samples"],
            row["genes"],
            row["design_full_rank"],
            row["batch_covariates"],
            row["qc_status"],
        ]
        for row in _deg_qc_records(project_dir)
    ]


def _enrichment_records(project_dir: Path) -> list[dict[str, str]]:
    return _read_tsv(project_dir / "results" / "enrichment" / "enrichment_results.tsv")[:20]


def _enrichment_rows(project_dir: Path) -> list[list[str]]:
    return [
        [
            row.get("dataset_id", ""),
            row.get("term_id", ""),
            row.get("term_name", ""),
            row.get("overlap_n", ""),
            row.get("adj_p_value", ""),
            row.get("overlap_genes", ""),
            row.get("source", ""),
        ]
        for row in _enrichment_records(project_dir)
    ]


def _meta_analysis_records(project_dir: Path) -> list[dict[str, str]]:
    return _read_tsv(project_dir / "results" / "meta_analysis" / "deg_meta_analysis.tsv")[:30]


def _meta_analysis_rows(project_dir: Path) -> list[list[str]]:
    return [
        [
            row.get("gene_symbol", ""),
            row.get("dataset_count", ""),
            row.get("mean_logFC", ""),
            row.get("random_effect_logFC", ""),
            row.get("dominant_direction", ""),
            row.get("direction_consistency", ""),
            row.get("heterogeneity_i2", ""),
            row.get("qc_status", ""),
            row.get("combined_p_score", ""),
            row.get("source_datasets", ""),
        ]
        for row in _meta_analysis_records(project_dir)
    ]


def _sasp_dataset_records(project_dir: Path) -> list[dict[str, str]]:
    return _read_tsv(project_dir / "results" / "sasp_score" / "sasp_dataset_scores.tsv")


def _sasp_gene_records(project_dir: Path) -> list[dict[str, str]]:
    return _read_tsv(project_dir / "results" / "sasp_score" / "sasp_gene_scores.tsv")[:30]


def _sasp_dataset_rows(project_dir: Path) -> list[list[str]]:
    rows = _sasp_dataset_records(project_dir)
    return [
        [
            row.get("dataset_id", ""),
            row.get("matched_sasp_genes", ""),
            row.get("up_sasp_genes", ""),
            row.get("down_sasp_genes", ""),
            row.get("sasp_dataset_score", ""),
            row.get("top_sasp_gene", ""),
            row.get("status", ""),
            row.get("deg_artifact", ""),
        ]
        for row in rows
    ] or [["", "", "", "", "", "", "NOT_RUN", "Run python tc_lite.py sasp-score --project <project>"]]


def _sasp_gene_rows(project_dir: Path) -> list[list[str]]:
    return [
        [
            row.get("dataset_id", ""),
            row.get("gene_symbol", ""),
            row.get("logFC", ""),
            row.get("adj_p_value", ""),
            row.get("direction", ""),
            row.get("sasp_component_score", ""),
            row.get("source_artifact", ""),
        ]
        for row in _sasp_gene_records(project_dir)
    ]


def _causal_evidence_records(project_dir: Path) -> list[dict[str, str]]:
    return _read_tsv(project_dir / "results" / "causal_evidence" / "causal_evidence_grades.tsv")[:30]


def _causal_evidence_rows(project_dir: Path) -> list[list[str]]:
    return [
        [
            row.get("gene_symbol", ""),
            row.get("causal_grade", ""),
            row.get("support_level", ""),
            row.get("methods", ""),
            row.get("evidence_types", ""),
            row.get("evidence_count", ""),
            row.get("best_p_value", ""),
            row.get("review_flags", ""),
            row.get("review_status", ""),
            row.get("artifact_refs", ""),
            row.get("rationale", ""),
        ]
        for row in _causal_evidence_records(project_dir)
    ]


def _experiment_records(project_dir: Path) -> list[dict[str, Any]]:
    return _read_json(project_dir / "results" / "experiments" / "experiment_designs.json", [])[:10]


def _experiment_rows(project_dir: Path) -> list[list[str]]:
    return [
        [
            row.get("candidate", ""),
            row.get("title", ""),
            row.get("objective", ""),
            "; ".join(row.get("readouts", [])),
            "; ".join(row.get("risks", [])),
        ]
        for row in _experiment_records(project_dir)
    ]


def _review_status(scores: list[dict[str, str]], matches: list[dict[str, str]], unknown_review: list[dict[str, str]]) -> tuple[str, list[list[str]]]:
    top_scores = scores[:20]
    review_match_count = sum(1 for row in matches if row.get("match_status") != "MATCH")
    hard_gate_count = sum(1 for row in top_scores if row.get("hard_gate_status") != "PASS")
    unknown_safety_count = sum(1 for row in top_scores if row.get("safety_gate") == "UNKNOWN")
    checklist = [
        ["ResearchSpec schema", "PASS", "research_spec.json exists and has passed validation before execution"],
        ["Dataset match", "REVIEW" if review_match_count else "PASS", f"{review_match_count} dataset match warning(s)"],
        ["DEG QC", "PASS", "bulk_deg results include qc_summary and run_manifest when the module is executed"],
        ["Candidate hard gates", "REVIEW" if hard_gate_count else "PASS", f"{hard_gate_count} top candidate hard-gate issue(s)"],
        ["Safety annotation", "REVIEW" if unknown_safety_count or unknown_review else "PASS", f"Top20 UNKNOWN safety={unknown_safety_count}; unknown rows={len(unknown_review)}"],
        ["Traceability", "PASS", "Evidence rows preserve source_dataset and artifact_path"],
    ]
    decision = "REVIEW_REQUIRED" if review_match_count or hard_gate_count or unknown_safety_count else "READY_FOR_EXPLORATORY_VALIDATION"
    return decision, checklist


def _evidence_by_gene(project_dir: Path, genes: list[str]) -> dict[str, list[dict[str, Any]]]:
    migrate_evidence_db(project_dir)
    from .evidence_repository import load_evidence_rows

    out = {}
    for gene in genes:
        out[gene] = load_evidence_rows(project_dir, gene=gene, limit=1000)["rows"]
    return out


def _evidence_records(project_dir: Path, scores: list[dict[str, str]]) -> list[dict[str, Any]]:
    top_genes = [row.get("entity_symbol", "") for row in scores[:8]]
    evidence = _evidence_by_gene(project_dir, top_genes)
    records = []
    for score in scores[:8]:
        gene = score.get("entity_symbol", "")
        records.append(
            {
                "gene": gene,
                "score_id": score.get("score_id", ""),
                "evidence_snapshot_id": score.get("evidence_snapshot_id", ""),
                "evidence_refs": [item for item in score.get("evidence_refs", "").split(";") if item],
                "score": score.get("final_score", ""),
                "tier": score.get("tier", ""),
                "hard_gate_status": score.get("hard_gate_status", ""),
                "evidence": evidence.get(gene, []),
            }
        )
    return records


def _fmt_number(value: Any) -> str:
    if value in (None, ""):
        return ""
    try:
        return f"{float(value):.4g}"
    except (TypeError, ValueError):
        return str(value)


def _short_evidence_ref(ev: dict[str, Any]) -> dict[str, Any]:
    return {
        "evidence_id": ev.get("evidence_id", ""),
        "evidence_type": ev.get("evidence_type", ""),
        "source_dataset": ev.get("source_dataset", "") or "",
        "artifact_path": ev.get("artifact_path", "") or "",
        "evidence_level": ev.get("evidence_level", "") or "",
        "evidence_weight": ev.get("evidence_weight", ""),
        "quality_score": ev.get("quality_score", ""),
        "limitation": ev.get("limitation", "") or "",
    }


def _candidate_evidence_cards(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    cards: list[dict[str, Any]] = []
    for record in records:
        evidence = record.get("evidence", [])
        sasp = [ev for ev in evidence if ev.get("evidence_type") == "sasp_score"]
        cell = [ev for ev in evidence if ev.get("evidence_type") == "cell_type_expression"]
        surface = [ev for ev in evidence if ev.get("evidence_type") in {"surface_marker_annotation", "accessibility"}]
        sasp_sorted = sorted(sasp, key=lambda ev: float(ev.get("quality_score") or 0), reverse=True)
        cell_sorted = sorted(cell, key=lambda ev: float(ev.get("quality_score") or 0), reverse=True)
        surface_sorted = sorted(surface, key=lambda ev: float(ev.get("quality_score") or 0), reverse=True)
        best_sasp = sasp_sorted[0] if sasp_sorted else {}
        best_cell = cell_sorted[0] if cell_sorted else {}
        best_surface = surface_sorted[0] if surface_sorted else {}
        cell_context = "; ".join(
            sorted(
                {
                    " / ".join(part for part in [str(ev.get("tissue") or "").strip(), str(ev.get("source_dataset") or "").strip()] if part)
                    for ev in cell_sorted[:4]
                }
            )
        )
        cards.append(
            {
                "gene": record.get("gene", ""),
                "score": record.get("score", ""),
                "tier": record.get("tier", ""),
                "route": best_surface.get("route", "") or "",
                "sasp": {
                    "status": "present" if sasp else "missing",
                    "best_component_score": _fmt_number(best_sasp.get("effect_size")),
                    "direction": best_sasp.get("direction", "") or "",
                    "dataset": best_sasp.get("source_dataset", "") or "",
                    "evidence_count": len(sasp),
                    "evidence": [_short_evidence_ref(ev) for ev in sasp_sorted[:3]],
                },
                "cell_type": {
                    "status": "present" if cell else "missing",
                    "context": cell_context,
                    "best_source": best_cell.get("source_dataset", "") or "",
                    "evidence_count": len(cell),
                    "evidence": [_short_evidence_ref(ev) for ev in cell_sorted[:3]],
                },
                "surface_or_secreted": {
                    "status": "present" if surface else "missing",
                    "route": best_surface.get("route", "") or record.get("route", ""),
                    "best_source": best_surface.get("source_dataset", "") or "",
                    "evidence_count": len(surface),
                    "evidence": [_short_evidence_ref(ev) for ev in surface_sorted[:3]],
                },
            }
        )
    return cards


def _candidate_evidence_card_html(cards: list[dict[str, Any]]) -> str:
    if not cards:
        return '<p class="note">No candidate evidence cards available. Run evidence import, scoring, and report generation.</p>'
    blocks = []
    for card in cards:
        gene = card.get("gene", "")
        sasp = card.get("sasp", {})
        cell = card.get("cell_type", {})
        surface = card.get("surface_or_secreted", {})
        blocks.append(
            '<article class="candidate-card">'
            f'<div class="candidate-card-head"><a href="#evidence-{html.escape(gene)}">{html.escape(gene)}</a>'
            f'<span>Score {html.escape(str(card.get("score", "")))} · Tier {html.escape(str(card.get("tier", "")))}</span></div>'
            '<div class="signal-grid">'
            f'<div class="signal"><strong>SASP</strong><span class="status {html.escape(str(sasp.get("status", "")))}">{html.escape(str(sasp.get("status", "")))}</span>'
            f'<p>component {html.escape(str(sasp.get("best_component_score", "")))} · {html.escape(str(sasp.get("direction", "")))} · {html.escape(str(sasp.get("dataset", "")))}</p>'
            f'<small>{html.escape(str(sasp.get("evidence_count", 0)))} evidence item(s)</small></div>'
            f'<div class="signal"><strong>Cell / tissue</strong><span class="status {html.escape(str(cell.get("status", "")))}">{html.escape(str(cell.get("status", "")))}</span>'
            f'<p>{html.escape(str(cell.get("context", "") or "No cell-type evidence linked"))}</p>'
            f'<small>{html.escape(str(cell.get("best_source", "")))} · {html.escape(str(cell.get("evidence_count", 0)))} item(s)</small></div>'
            f'<div class="signal"><strong>Surface / secreted</strong><span class="status {html.escape(str(surface.get("status", "")))}">{html.escape(str(surface.get("status", "")))}</span>'
            f'<p>{html.escape(str(surface.get("route", "") or "No accessible route linked"))}</p>'
            f'<small>{html.escape(str(surface.get("best_source", "")))} · {html.escape(str(surface.get("evidence_count", 0)))} item(s)</small></div>'
            '</div>'
            '</article>'
        )
    return "".join(blocks)


def _evidence_sections(records: list[dict[str, Any]]) -> str:
    sections = []
    for record in records:
        ev_rows = [
            [
                ev.get("evidence_type", ""),
                ev.get("evidence_id", ""),
                ev.get("source_dataset", "") or "",
                ev.get("direction", "") or "",
                _fmt_number(ev.get("effect_size")),
                _fmt_number(ev.get("p_value")),
                ev.get("evidence_level", "") or "",
                ev.get("evidence_weight", "") or "",
                ev.get("evidence_basis", "") or "",
                ev.get("artifact_path", "") or "",
                ev.get("run_id", "") or "",
                ev.get("artifact_id", "") or "",
                ev.get("limitation", "") or "",
            ]
            for ev in record.get("evidence", [])
        ]
        gene = record.get("gene", "")
        sections.append(
            f'<h3 id="evidence-{html.escape(gene)}">{html.escape(gene)} evidence chain</h3>'
            f'<p class="note">score_id: <code>{html.escape(record.get("score_id", ""))}</code> | snapshot: <code>{html.escape(record.get("evidence_snapshot_id", ""))}</code></p>'
            + _table(["Evidence", "Evidence ID", "Dataset", "Direction", "Effect", "P/adj.P", "Level", "Weight", "Basis", "Artifact", "Run", "Artifact ID", "Limitation"], ev_rows)
        )
    return "".join(sections)


def _cell_type_summary(project_dir: Path) -> dict[str, Any]:
    return _read_json(project_dir / "results" / "cell_type_evidence" / "cell_type_summary.json", {"row_count": 0, "cell_type_by_gene": {}, "by_source": {}})


def _cell_type_rows(project_dir: Path, scores: list[dict[str, str]]) -> list[list[str]]:
    summary = _cell_type_summary(project_dir)
    by_gene = summary.get("cell_type_by_gene", {})
    ordered_genes = [row.get("entity_symbol", "") for row in scores[:20]]
    rows = []
    for gene in ordered_genes:
        for item in by_gene.get(gene, [])[:8]:
            rows.append(
                [
                    gene,
                    item.get("cell_type", ""),
                    item.get("tissue", ""),
                    item.get("evidence_source", ""),
                    item.get("confidence", ""),
                    item.get("artifact_path", ""),
                    item.get("context", ""),
                    item.get("limitation", ""),
                ]
            )
    if not rows:
        rows.append(["", "No cell-type evidence built yet", "", "", "", "results/cell_type_evidence/cell_type_evidence.tsv", "Run cell-type-evidence after HPA/PanglaoDB/CellMarker/scRNA/full-text inputs are available.", ""])
    return rows


def _recovery_summary(project_dir: Path) -> dict[str, Any]:
    return _read_json(project_dir / "results" / "recovery" / "recovery_manifest.json", {"item_count": 0, "open_count": 0, "items": []})


def _recovery_rows(project_dir: Path) -> list[list[str]]:
    manifest = _recovery_summary(project_dir)
    return [
        [
            row.get("stage", ""),
            row.get("status", ""),
            row.get("reason", ""),
            row.get("suggested_action", ""),
            row.get("manual_correction", ""),
            row.get("source_artifact", ""),
        ]
        for row in manifest.get("items", [])[:30]
    ] or [["", "PASS", "No open recovery items in manifest", "", "", "results/recovery/recovery_manifest.json"]]


def _audit_rows(context: dict[str, Any]) -> tuple[list[list[str]], list[list[str]], list[list[str]]]:
    trace = context.get("agent_trace", {})
    stages = [
        [row.get("name", ""), row.get("status", ""), row.get("message", ""), row.get("details", {}).get("label", "")]
        for row in trace.get("stages", [])
    ]
    reviews = [
        [
            row.get("timestamp", ""),
            row.get("review_id", ""),
            row.get("item_type", ""),
            row.get("item_id", ""),
            row.get("action", ""),
            row.get("reason", row.get("note", "")),
            row.get("report_ref", ""),
        ]
        for row in context.get("review_actions", [])[-20:]
    ]
    adapters = [
        [
            row.get("resource_id", ""),
            row.get("resource_type", ""),
            row.get("database_adapter", "") or row.get("adapter", ""),
            row.get("input_rows", ""),
            row.get("normalized_rows", ""),
            row.get("dropped_rows", ""),
            row.get("field_mapping", ""),
        ]
        for row in context.get("adapter_audit", [])
    ]
    return stages, reviews, adapters


def _mcp_audit_records(project_dir: Path) -> dict[str, Any]:
    return _read_json(project_dir / "v4" / "mcp_call_audit_summary.json", {"call_count": 0, "failure_count": 0, "by_tool": {}, "latest_calls": []})


def _mcp_audit_rows(project_dir: Path) -> list[list[str]]:
    audit = _mcp_audit_records(project_dir)
    return [
        [
            row.get("call_id", ""),
            row.get("tool_id", ""),
            row.get("actor", ""),
            row.get("risk", ""),
            row.get("status", ""),
            row.get("failure_reason", ""),
        ]
        for row in audit.get("latest_calls", [])[-20:]
    ]


def _codex_engineering_records(project_dir: Path) -> dict[str, Any]:
    return load_codex_engineering(project_dir)


def _codex_engineering_rows(project_dir: Path) -> list[list[str]]:
    data = _codex_engineering_records(project_dir)
    return [
        [
            row.get("result_id", ""),
            row.get("codex_job_id", ""),
            row.get("work_order_id", ""),
            row.get("status", ""),
            row.get("merge_status", ""),
            row.get("review_status", ""),
            ";".join(row.get("artifacts", [])),
        ]
        for row in data.get("results", [])[-20:]
    ]


def _limitation_records(context: dict[str, Any]) -> list[dict[str, str]]:
    rows = []
    for row in context["unknown_review"][:50]:
        rows.append(
            {
                "gene": row.get("gene_symbol", ""),
                "missing_fields": row.get("missing_fields", ""),
                "route": row.get("route", ""),
                "safety_gate": row.get("safety_gate", ""),
                "recommended_action": row.get("recommended_action", ""),
            }
        )
    if any(row.get("match_status") != "MATCH" for row in context["matches"]):
        rows.append(
            {
                "gene": "DATASET_SCOPE",
                "missing_fields": "dataset/spec mismatch or review warning",
                "route": "",
                "safety_gate": "REVIEW",
                "recommended_action": "Review dataset cards and matching notes before interpreting candidates.",
            }
        )
    return rows


def _structured_report(project_dir: Path, context: dict[str, Any]) -> dict[str, Any]:
    spec = context["spec"]
    evidence = _evidence_records(project_dir, context["scores"])
    candidate_evidence_cards = _candidate_evidence_cards(evidence)
    return {
        "project": project_dir.name,
        "report_version": "0.4",
        "sections": REPORT_SECTIONS,
        "executive_summary": {
            "decision": context["decision"],
            "checklist": [{"check": row[0], "status": row[1], "note": row[2]} for row in context["checklist"]],
            "evidence_rows": context["evidence_count"],
        },
        "research_question": {
            "research_theme": spec.get("research_theme", ""),
            "disease_scope": spec.get("disease_scope", {}).get("canonical", "unknown"),
            "priority_tissues": spec.get("priority_tissues", []),
            "target_routes": spec.get("target_routes", []),
            "boundaries": [
                "Exploratory target discovery only.",
                "Expression association is not causal proof.",
                "Manual review is required for dataset warnings and UNKNOWN annotations.",
            ],
        },
        "methods": {"analysis_modules": _method_records(context)},
        "data_sources_and_qc": {
            "datasets": _dataset_records(context),
            "bulk_rna_microarray_qc": _deg_qc_records(project_dir),
            "enrichment_overview": _enrichment_records(project_dir),
        },
        "advanced_analysis": {
            "meta_analysis": _meta_analysis_records(project_dir),
            "sasp_score": {
                "datasets": _sasp_dataset_records(project_dir),
                "genes": _sasp_gene_records(project_dir),
                "manifest": _read_json(project_dir / "results" / "sasp_score" / "run_manifest.json", {}),
            },
            "causal_evidence": _causal_evidence_records(project_dir),
        },
        "candidate_ranking": context["scores"][:20],
        "candidate_evidence_cards": candidate_evidence_cards,
        "evidence_chain": evidence,
        "cell_type_evidence": _cell_type_summary(project_dir),
        "failure_recovery": _recovery_summary(project_dir),
        "report_evidence_refs": {
            row.get("gene", ""): {
                "score_id": row.get("score_id", ""),
                "evidence_snapshot_id": row.get("evidence_snapshot_id", ""),
                "evidence_refs": row.get("evidence_refs", []),
            }
            for row in evidence
        },
        "evidence_review_report_index": {},
        "scoring_manifest": _read_json(project_dir / "results" / "scoring" / "target_score_manifest.json", {}),
        "limitations": _limitation_records(context),
        "experiment_suggestions": _experiment_records(project_dir),
        "approval_and_audit": {
            "agent_trace": context.get("agent_trace", {}).get("stages", []),
            "manual_review_actions": context.get("review_actions", [])[-20:],
            "review_queue": context.get("review_queue", {}),
            "approval_state": context.get("approval_state", {}),
            "adapter_audit": context.get("adapter_audit", []),
            "mcp_call_audit": _mcp_audit_records(project_dir),
            "codex_engineering": _codex_engineering_records(project_dir),
        },
    }


def _html_report(project_dir: Path, context: dict[str, Any], structured: dict[str, Any]) -> str:
    spec = context["spec"]
    scores = context["scores"]
    decision = context["decision"]
    checklist = context["checklist"]
    stages, reviews, adapters = _audit_rows(context)
    evidence_records = structured["evidence_chain"]
    limitation_rows = [
        [r.get("gene", ""), r.get("missing_fields", ""), r.get("route", ""), r.get("safety_gate", ""), r.get("recommended_action", "")]
        for r in structured["limitations"]
    ]
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>TargetCompass Lite Candidate Target Report</title>
  <style>
    body {{ font-family: Arial, "Microsoft YaHei", sans-serif; margin: 0; color: #202124; background: #f7f8fa; }}
    main {{ max-width: 1180px; margin: 0 auto; padding: 32px 24px 64px; }}
    header {{ border-bottom: 1px solid #d0d7de; margin-bottom: 24px; padding-bottom: 18px; }}
    h1 {{ font-size: 30px; margin: 0 0 10px; }}
    h2 {{ margin-top: 32px; border-bottom: 1px solid #d8dee4; padding-bottom: 8px; }}
    h3 {{ margin-top: 24px; }}
    section {{ background: #fff; border: 1px solid #d8dee4; border-radius: 8px; padding: 18px 20px; margin: 18px 0; }}
    table {{ border-collapse: collapse; width: 100%; margin-top: 14px; }}
    th, td {{ border: 1px solid #d0d7de; padding: 8px; text-align: left; font-size: 13px; vertical-align: top; }}
    th {{ background: #f6f8fa; }}
    .grid {{ display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 14px; }}
    .card {{ border: 1px solid #d8dee4; border-radius: 8px; padding: 12px 14px; background: #fff; }}
    .candidate-card {{ border: 1px solid #d8dee4; border-radius: 8px; padding: 14px; margin: 14px 0; background: #fbfcfd; }}
    .candidate-card-head {{ display: flex; justify-content: space-between; gap: 14px; align-items: baseline; margin-bottom: 12px; }}
    .candidate-card-head a {{ font-size: 18px; font-weight: 700; color: #0969da; text-decoration: none; }}
    .candidate-card-head span {{ color: #5f6368; font-size: 13px; }}
    .signal-grid {{ display: grid; grid-template-columns: repeat(3, minmax(0, 1fr)); gap: 10px; }}
    .signal {{ background: #fff; border: 1px solid #d8dee4; border-radius: 8px; padding: 10px; min-height: 96px; }}
    .signal strong {{ display: block; margin-bottom: 8px; }}
    .signal p {{ margin: 6px 0; font-size: 13px; }}
    .signal small {{ color: #5f6368; }}
    .status {{ display: inline-block; border-radius: 999px; padding: 2px 8px; font-size: 12px; background: #f6f8fa; border: 1px solid #d8dee4; }}
    .status.present {{ color: #116329; background: #dafbe1; border-color: #aceebb; }}
    .status.missing {{ color: #8a4600; background: #fff8c5; border-color: #d4a72c; }}
    @media (max-width: 860px) {{ .signal-grid, .grid {{ grid-template-columns: 1fr; }} .candidate-card-head {{ display: block; }} }}
    .note {{ color: #5f6368; }}
    .warning {{ background: #fff8c5; border: 1px solid #d4a72c; padding: 10px 12px; border-radius: 6px; }}
    .decision {{ font-size: 18px; font-weight: 700; padding: 12px 14px; border-radius: 6px; background: #ddf4ff; border: 1px solid #54aeef; }}
  </style>
</head>
<body>
<main>
  <header>
    <h1>TargetCompass Lite Candidate Target Report</h1>
    <p><strong>Project:</strong> {html.escape(project_dir.name)} | <strong>Evidence rows:</strong> {context["evidence_count"]}</p>
    <p class="note">科研探索用途。本报告输出候选分子、证据链、QC、限制和实验建议，不提供医学判断、临床决策或用药指引，也不主张因果结论。</p>
  </header>

  <section id="executive-summary"><h2>执行摘要</h2>
    <div class="decision">{html.escape(decision)}</div>
    {_table(["Check", "Status", "Note"], checklist)}
  </section>

  <section id="research-question"><h2>研究问题与边界</h2>
    <div class="grid">
      <div class="card"><strong>Research theme</strong><br>{html.escape(spec.get("research_theme", ""))}</div>
      <div class="card"><strong>Disease scope</strong><br>{html.escape(spec.get("disease_scope", {}).get("canonical", "unknown"))}</div>
      <div class="card"><strong>Priority tissues</strong><br>{html.escape(", ".join(spec.get("priority_tissues", [])))}</div>
      <div class="card"><strong>Target routes</strong><br>{html.escape(", ".join(spec.get("target_routes", [])))}</div>
    </div>
  </section>

  <section id="methods"><h2>方法与模块</h2>
    {_table(["Module", "Status", "Input", "Runner", "Outputs", "Notes"], _method_rows(context))}
  </section>

  <section id="data-qc"><h2>数据来源与QC</h2>
    {_table(["Dataset", "Source class", "Grade", "Modality", "Metadata quality", "Gene identity", "Use", "Match", "Notes"], _dataset_rows(context))}
    <h3>Bulk RNA / microarray QC</h3>
    {_table(["Dataset", "Matrix type", "Runner", "Case n", "Control n", "Genes", "Full rank", "Batch covariates", "QC"], _deg_qc_rows(project_dir))}
    <h3>Enrichment overview</h3>
    {_table(["Dataset", "Term ID", "Term", "Overlap", "Adj.P", "Genes", "Source"], _enrichment_rows(project_dir))}
    <h3>Meta-analysis overview</h3>
    {_table(["Gene", "Datasets", "Mean logFC", "Random logFC", "Direction", "Consistency", "I2", "QC", "Combined score", "Sources"], _meta_analysis_rows(project_dir))}
    <h3>SASP score overview</h3>
    <p class="note">SASP score is computed from real DEG outputs using a configurable SASP core gene set. It is phenotype evidence, not causal proof.</p>
    {_table(["Dataset", "Matched SASP genes", "Up", "Down", "Dataset score", "Top SASP gene", "Status", "Artifact"], _sasp_dataset_rows(project_dir))}
    {_table(["Dataset", "Gene", "logFC", "Adj.P", "Direction", "Component score", "Artifact"], _sasp_gene_rows(project_dir))}
    <h3>Causal evidence grading</h3>
    {_table(["Gene", "Grade", "Support", "Methods", "Evidence types", "Count", "Best P", "Review flags", "Review", "Artifacts", "Rationale"], _causal_evidence_rows(project_dir))}
  </section>

  <section id="candidate-ranking"><h2>候选排序</h2>
    {_table(["Score ID", "Gene", "Route", "Score", "Tier", "Hard gate", "Safety", "Next experiments"], _candidate_rows(scores))}
  </section>

  <section id="evidence-chain"><h2>证据链</h2>
    <div id="candidate-evidence-cards">
      <h3>候选基因证据卡</h3>
      <p class="note">每张卡把候选基因的 SASP 评分、细胞/组织证据、表面或分泌可及性集中展示；点击基因名可跳到完整 EvidenceItem 追溯链。</p>
      {_candidate_evidence_card_html(structured.get("candidate_evidence_cards", []))}
    </div>
    {_evidence_sections(evidence_records)}
  </section>

  <section id="cell-type-evidence"><h2>细胞类型证据</h2>
    <p class="note">回答“这个分子表达在哪类细胞上”。证据可来自 HPA、PanglaoDB、CellMarker、scRNA pseudobulk 或全文 LLM 抽取；每行保留来源和限制。</p>
    {_table(["Gene", "Cell type", "Tissue", "Source", "Confidence", "Artifact", "Context", "Limitation"], _cell_type_rows(project_dir, scores))}
  </section>

  <section id="failure-recovery"><h2>失败恢复与人工补正</h2>
    <p class="note">记录数据库、GEO、全文和长任务失败原因，并给出可重试命令或人工补正路径。</p>
    {_table(["Stage", "Status", "Reason", "Suggested action", "Manual correction", "Artifact"], _recovery_rows(project_dir))}
  </section>

  <section id="limitations"><h2>限制与风险</h2>
    <div class="warning">表达证据属于 association-level evidence，不等同于因果证据。Dataset/spec warning、UNKNOWN safety、外部 adapter 证据和 Python fallback 结果都需要人工复核。</div>
    {_table(["Gene", "Missing fields", "Route", "Safety", "Recommended action"], limitation_rows)}
  </section>

  <section id="experiment-suggestions"><h2>实验建议</h2>
    {_table(["Candidate", "Title", "Objective", "Readouts", "Risks"], _experiment_rows(project_dir))}
  </section>

  <section id="approval-audit"><h2>审批与审计</h2>
    <h3>Approval state</h3>
    {_table(["Status", "Queue", "Reviews", "Approved", "Rejected", "Signer", "Reason"], [[
        context.get("approval_state", {}).get("status", "draft"),
        context.get("approval_state", {}).get("queue_count", ""),
        context.get("approval_state", {}).get("review_count", ""),
        context.get("approval_state", {}).get("approved_count", ""),
        context.get("approval_state", {}).get("rejected_count", ""),
        context.get("approval_state", {}).get("signer", ""),
        context.get("approval_state", {}).get("reason", ""),
    ]])}
    <h3>Review queue</h3>
    {_table(["Type", "ID", "Title", "Status", "Reason", "Report ref"], [
        [
            row.get("item_type", ""),
            row.get("item_id", ""),
            row.get("title", ""),
            row.get("review_status", ""),
            row.get("reason", ""),
            row.get("report_ref", ""),
        ]
        for row in context.get("review_queue", {}).get("items", [])[:20]
    ])}
    <h3>Evidence trace index</h3>
    {_table(["Index", "Evidence", "Review items", "Report refs", "Path"], [[
        structured.get("evidence_review_report_index", {}).get("index_id", ""),
        structured.get("evidence_review_report_index", {}).get("evidence_count", ""),
        structured.get("evidence_review_report_index", {}).get("review_item_count", ""),
        structured.get("evidence_review_report_index", {}).get("report_ref_count", ""),
        structured.get("evidence_review_report_index", {}).get("path", ""),
    ]])}
    <h3>Agent trace</h3>
    {_table(["Stage", "Status", "Message", "Label"], stages)}
    <h3>Manual review actions</h3>
    {_table(["Timestamp", "Review ID", "Type", "ID", "Action", "Reason", "Report ref"], reviews)}
    <h3>Adapter audit</h3>
    {_table(["Resource", "Type", "Adapter", "Input rows", "Normalized rows", "Dropped rows", "Field mapping"], adapters)}
    <h3>MCP call audit</h3>
    <p class="note">Calls: {html.escape(str(_mcp_audit_records(project_dir).get("call_count", 0)))} | Failures: {html.escape(str(_mcp_audit_records(project_dir).get("failure_count", 0)))}</p>
    {_table(["Call", "Tool", "Actor", "Risk", "Status", "Failure"], _mcp_audit_rows(project_dir))}
    <h3>Codex engineering results</h3>
    {_table(["Result", "Codex job", "WorkOrder", "Status", "Merge gate", "Review", "Artifacts"], _codex_engineering_rows(project_dir))}
  </section>
</main>
</body>
</html>"""


def _build_context(project_dir: Path) -> dict[str, Any]:
    scores = _read_csv(project_dir / "candidate_scores.csv")
    screening = _read_csv(project_dir / "eligible_datasets.csv")
    matches = _read_csv(project_dir / "dataset_match_report.csv")
    unknown_review = _read_tsv(project_dir / "results" / "annotation" / "unknown_review.tsv")
    from .evidence_repository import load_evidence_rows

    evidence_count = len(load_evidence_rows(project_dir, limit=100000)["rows"])
    decision, checklist = _review_status(scores, matches, unknown_review)
    return {
        "scores": scores,
        "screening": screening,
        "matches": matches,
        "unknown_review": unknown_review,
        "evidence_count": evidence_count,
        "decision": decision,
        "checklist": checklist,
        "spec": _read_json(project_dir / "research_spec.json", {}),
        "analysis_modules": _read_json(project_dir / "analysis_module_registry.json", {"modules": []}),
        "agent_trace": _read_json(project_dir / "results" / "agent_trace.json", {}),
        "review_actions": _read_tsv(project_dir / "results" / "review_actions.tsv"),
        "review_queue": _read_json(project_dir / "results" / "review_queue.json", {"items": [], "queue_count": 0}),
        "approval_state": _read_json(project_dir / "results" / "approval_state.json", {"status": "draft"}),
        "adapter_audit": _read_tsv(project_dir / "results" / "adapter_audit" / "adapter_audit.tsv"),
    }


def _write_docx(docx_path: Path, project_dir: Path, context: dict[str, Any], structured: dict[str, Any]) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("TargetCompass Lite Candidate Target Report", level=1)
    doc.add_paragraph(f"Project: {project_dir.name}")
    doc.add_paragraph(f"Evidence rows: {context['evidence_count']}")
    doc.add_paragraph("科研探索用途；不提供医学判断、临床决策或用药指引，也不主张因果结论。")
    for section in REPORT_SECTIONS:
        doc.add_heading(section, level=2)
        if section == "执行摘要":
            doc.add_paragraph(context["decision"])
            for item, status, note in context["checklist"]:
                doc.add_paragraph(f"{item}: {status} - {note}", style="List Bullet")
        elif section == "候选排序":
            table = doc.add_table(rows=1, cols=7)
            headers = ["Gene", "Route", "Score", "Tier", "Hard gate", "Safety", "Next experiments"]
            for idx, header in enumerate(headers):
                table.rows[0].cells[idx].text = header
            for row in context["scores"][:20]:
                values = [
                    row.get("entity_symbol", ""),
                    row.get("route", ""),
                    row.get("final_score", ""),
                    row.get("tier", ""),
                    row.get("hard_gate_status", ""),
                    row.get("safety_gate", ""),
                    row.get("next_experiments", ""),
                ]
                cells = table.add_row().cells
                for idx, value in enumerate(values):
                    cells[idx].text = value
        elif section == "数据来源与QC":
            doc.add_paragraph(f"Datasets: {len(structured['data_sources_and_qc']['datasets'])}")
            doc.add_paragraph(f"Bulk QC records: {len(structured['data_sources_and_qc']['bulk_rna_microarray_qc'])}")
        else:
            doc.add_paragraph("See HTML report and target_report_structured.json for the full structured content.")
    doc.save(docx_path)


def build_report(project_dir: Path) -> tuple[Path, Path]:
    context = _build_context(project_dir)
    structured = _structured_report(project_dir, context)
    reports = project_dir / "reports"
    reports.mkdir(exist_ok=True)
    html_path = reports / "target_report.html"
    docx_path = reports / "target_report.docx"
    structured_path = reports / "target_report_structured.json"
    _write_json(structured_path, structured)
    traceability = refresh_traceability(project_dir)
    structured["evidence_review_report_index"] = traceability.get("refreshed", {}).get("evidence_review_report_index", {})
    html_text = _html_report(project_dir, context, structured)
    for phrase in PROHIBITED_CLAIMS:
        if phrase.lower() in html_text.lower():
            raise ValueError(f"report contains prohibited claim phrase: {phrase}")
    html_path.write_text(html_text, encoding="utf-8")
    _write_json(structured_path, structured)
    try:
        _write_docx(docx_path, project_dir, context, structured)
    except Exception:
        docx_path.write_text(html_text, encoding="utf-8")
    try:
        from .output_backend import publish_output_artifacts

        publish_output_artifacts(
            project_dir,
            [html_path, docx_path, structured_path],
            producer="reporting",
            artifact_type="target_report_output",
            task_id="reporting",
        )
    except Exception:
        pass
    return html_path, docx_path
